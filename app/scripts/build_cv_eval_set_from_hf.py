import json,re
from pathlib import Path
from datetime import datetime
from docx import Document

from app.ai.skill_canonicalization import canonicalize_skill

SRC = Path("data/raw/resumes_normalized.jsonl")
DOC_DIR = Path("data/cv_hf_docs")
LBL = Path("data/labels/cv_extraction_hf_labels.jsonl")
MAX_ROWS = 800
MIN_SKILLS = 3

UNKNOWN = {"", "unknown", "n/a", "na", "none", "null", "-", "--", "not provided"}
LEVEL_WORDS = {"beginner","intermediate","advanced","expert","native","fluent","unknown"}
_DURATION_SKILL_RE = re.compile(r"^\d+(?:\.\d+)?\s*(?:months?|years?|yrs?)(?:\s+experience)?$", re.I)

def clean(v):
    if v is None: return ""
    s = str(v).strip()
    return "" if s.lower() in UNKNOWN else s

def walk_skill_names(node):
    if isinstance(node, dict):
        if "name" in node:
            n = clean(node.get("name"))
            if n:
                yield n
        for v in node.values():
            yield from walk_skill_names(v)
    elif isinstance(node, list):
        for x in node:
            yield from walk_skill_names(x)
    elif isinstance(node, str):
        s = clean(node)
        low = s.lower()
        if s and low not in LEVEL_WORDS and len(s.split()) <= 4:
            yield s

def norm_skill(s):
    normalized = canonicalize_skill(s)
    if not normalized:
        return ""
    if _DURATION_SKILL_RE.fullmatch(normalized):
        return ""
    if re.fullmatch(r"[a-z0-9]+/", normalized):
        return ""
    return normalized

def uniq(seq):
    seen=set(); out=[]
    for x in seq:
        k=norm_skill(x)
        if k and k not in seen:
            seen.add(k); out.append(x.strip())
    return out

RANGE_RE = re.compile(r"(19\d{2}|20\d{2})\s*[-/]\s*(present|current|now|19\d{2}|20\d{2})", re.I)
YEARS_RE = re.compile(r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)", re.I)

def estimate_years(experience):
    cur = datetime.utcnow().year
    vals=[]
    for row in experience if isinstance(experience,list) else []:
        txt = " ".join([clean(row.get("year")), clean(row.get("details"))]) if isinstance(row,dict) else clean(row)
        for a,b in RANGE_RE.findall(txt):
            a=int(a); b=cur if b.lower() in {"present","current","now"} else int(b)
            if b>=a:
                vals.append(float(b-a))
        for y in YEARS_RE.findall(txt):
            vals.append(float(y))
    if not vals:
        return None
    return round(min(max(vals), 40.0),1)

def pick_title(personal_info, experience):
    if isinstance(experience,list):
        for row in experience:
            if isinstance(row,dict):
                t=clean(row.get("title"))
                if t:
                    return t
    if isinstance(personal_info,dict):
        summary = clean(personal_info.get("summary"))
        if summary:
            return " ".join(summary.split()[:8])
    return None

DOC_DIR.mkdir(parents=True, exist_ok=True)
LBL.parent.mkdir(parents=True, exist_ok=True)

written=0
with SRC.open("r",encoding="utf-8") as f, LBL.open("w",encoding="utf-8") as out:
    for i,line in enumerate(f, start=1):
        if written >= MAX_ROWS:
            break
        row=json.loads(line)
        pinfo=row.get("personal_info") or {}
        exp=row.get("experience") or []
        edu=row.get("education") or []
        skills=uniq(list(walk_skill_names(row.get("skills") or {})))
        title=pick_title(pinfo, exp)
        years=estimate_years(exp)

        if len(skills) < MIN_SKILLS and not title:
            continue

        doc = Document()
        name = clean((pinfo or {}).get("name")) or f"Candidate {i}"
        doc.add_paragraph(name)
        summary = clean((pinfo or {}).get("summary"))
        if summary:
            doc.add_paragraph(summary)

        if exp:
            doc.add_paragraph("Experience")
            for e in exp[:6]:
                if isinstance(e,dict):
                    line2 = " - ".join([x for x in [clean(e.get("title")), clean(e.get("year")), clean(e.get("details"))] if x])
                else:
                    line2 = clean(e)
                if line2:
                    doc.add_paragraph(line2)

        if edu:
            doc.add_paragraph("Education")
            for e in edu[:4]:
                if isinstance(e,dict):
                    line2 = " - ".join([x for x in [clean(e.get("title")), clean(e.get("year")), clean(e.get("details"))] if x])
                else:
                    line2 = clean(e)
                if line2:
                    doc.add_paragraph(line2)

        if skills:
            doc.add_paragraph("Skills")
            doc.add_paragraph(", ".join(skills[:60]))

        p = DOC_DIR / f"hf_{i:05d}.docx"
        doc.save(p)

        label = {
            "path": str(p.resolve()),
            "expected_skills": skills,
            "expected_title": title,
            "expected_experience_years": years,
        }
        out.write(json.dumps(label, ensure_ascii=False) + "\n")
        written += 1

print("written_docs=", written)
print("labels=", LBL)
print("docs_dir=", DOC_DIR)
